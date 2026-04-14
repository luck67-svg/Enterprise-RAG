from pathlib import Path
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_core.documents import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def load_file(path: str | Path) -> list[Document]:
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return PyMuPDFLoader(str(path)).load()

    if suffix == ".docx":
        return _load_docx(path)

    if suffix == ".txt":
        return _load_txt(path)

    raise ValueError(f"不支持的文件类型: {suffix}，支持: {', '.join(SUPPORTED_EXTENSIONS)}")


_DOCX_PAGE_CHARS = 1800  # 每"页"估算字符数（约 A4 一页中文）


def _load_docx(path: Path) -> list[Document]:
    """按段落切分 docx，估算页码，返回多个 Document。

    遇到 Heading 样式段落时强制落盘当前 buffer，使每章节成为独立 Document；
    单段超过 _DOCX_PAGE_CHARS 字符时也强制落盘，避免超大 chunk。
    page 字段按累计字符数估算，使来源引用比恒为 p.1 更有参考价值。
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    documents: list[Document] = []
    buffer: list[str] = []
    char_count = 0
    page = 1

    def _flush() -> None:
        nonlocal char_count, page
        if buffer:
            documents.append(Document(
                page_content="\n\n".join(buffer),
                metadata={"source": path.name, "page": page},
            ))
            page += max(1, char_count // _DOCX_PAGE_CHARS)
            buffer.clear()
            char_count = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # 遇到标题样式时先落盘当前 buffer，使标题成为新段的起点
        if para.style.name.startswith("Heading"):
            _flush()
        buffer.append(text)
        char_count += len(text)
        # 超过单页字符上限时强制落盘
        if char_count >= _DOCX_PAGE_CHARS:
            _flush()

    _flush()

    if not documents:
        # 空文件兜底
        documents.append(Document(page_content="", metadata={"source": path.name, "page": 1}))

    return documents


def _load_txt(path: Path) -> list[Document]:
    text = path.read_text(encoding="utf-8")
    return [Document(page_content=text, metadata={"source": path.name, "page": 1})]
